import os
import pypdf
from docx import Document


class ResumeExtractor:
    """Utility service to extract plain text from PDF, DOCX, and TXT files."""

    @staticmethod
    def extract_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def extract_pdf(file_path: str) -> str:
        reader = pypdf.PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n".join(text_parts)

    @staticmethod
    def extract_docx(file_path: str) -> str:
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)

    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """Determines the extension and extracts candidate resume content."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        _, ext = os.path.splitext(file_path.lower())
        if ext == ".pdf":
            return cls.extract_pdf(file_path)
        elif ext == ".docx":
            return cls.extract_docx(file_path)
        elif ext in (".txt", ".md"):
            return cls.extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported resume file extension: {ext}")
